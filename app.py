import streamlit as st
import time
import os
import json
import base64
import requests
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
from playwright.sync_api import sync_playwright

# =========================
# 도착지 추가 기능 - 경로/상수
# =========================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
MAP_DIR = os.path.join(BASE_DIR, "map")
CUSTOM_DEST_FILE = os.path.join(BASE_DIR, "custom_destinations.json")
ADD_DEST_OPTION = "➕ 도착지 추가하기"

# =========================
# GitHub 자동 커밋 설정
# =========================
GITHUB_OWNER = "myjeong14-cmyk"
GITHUB_REPO = "hrdk-report-generater"
GITHUB_BRANCH = "main"
GITHUB_API_BASE = f"https://api.github.com/repos/{GITHUB_OWNER}/{GITHUB_REPO}/contents"
GITHUB_RAW_BASE = f"https://raw.githubusercontent.com/{GITHUB_OWNER}/{GITHUB_REPO}/{GITHUB_BRANCH}"


def _github_token():
    """Streamlit secrets 또는 환경변수에서 GitHub 토큰을 읽는다."""
    try:
        if "GITHUB_TOKEN" in st.secrets:
            return st.secrets["GITHUB_TOKEN"]
    except Exception:
        pass
    return os.environ.get("GITHUB_TOKEN", "")


def _github_headers():
    token = _github_token()
    if not token:
        return None
    return {
        "Authorization": f"token {token}",
        "Accept": "application/vnd.github+json",
    }


def github_commit_file(path_in_repo, content_bytes, commit_message):
    """저장소의 path_in_repo 경로에 파일을 생성/갱신(commit)한다."""
    headers = _github_headers()
    if not headers:
        return False, "GITHUB_TOKEN이 설정되어 있지 않아 GitHub에 커밋할 수 없습니다."

    url = f"{GITHUB_API_BASE}/{path_in_repo}"
    try:
        get_resp = requests.get(url, headers=headers, params={"ref": GITHUB_BRANCH}, timeout=10)
        sha = get_resp.json().get("sha") if get_resp.status_code == 200 else None

        payload = {
            "message": commit_message,
            "content": base64.b64encode(content_bytes).decode("utf-8"),
            "branch": GITHUB_BRANCH,
        }
        if sha:
            payload["sha"] = sha

        put_resp = requests.put(url, headers=headers, json=payload, timeout=15)
        if put_resp.status_code in (200, 201):
            return True, None
        return False, f"GitHub 커밋 실패 ({put_resp.status_code}): {put_resp.text[:300]}"
    except Exception as e:
        return False, f"GitHub 커밋 중 오류: {e}"


def github_fetch_file(path_in_repo):
    """저장소에서 path_in_repo 경로의 파일 내용을 raw bytes로 가져온다. 실패 시 None."""
    try:
        resp = requests.get(f"{GITHUB_RAW_BASE}/{path_in_repo}", timeout=10)
        if resp.status_code == 200:
            return resp.content
    except Exception:
        pass
    return None


def load_custom_destinations():
    """사용자가 추가한 도착지 목록을 불러온다. GitHub 원본을 우선 조회하고,
    실패 시 로컬 캐시 파일을 사용한다."""
    remote_bytes = github_fetch_file("custom_destinations.json")
    if remote_bytes is not None:
        try:
            data = json.loads(remote_bytes.decode("utf-8"))
            # 로컬에도 캐시해서 같은 세션 내에서는 재조회 없이 사용 가능하게 함
            try:
                with open(CUSTOM_DEST_FILE, "w", encoding="utf-8") as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
            except Exception:
                pass
            return data
        except Exception:
            pass

    if os.path.exists(CUSTOM_DEST_FILE):
        try:
            with open(CUSTOM_DEST_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}
    return {}


def save_new_destination(name, region, dist):
    """새 도착지를 DESTINATION_DB 규칙에 맞게 계산하여 저장하고 GitHub에 커밋한다."""
    dist = float(dist)
    entry = {
        "dist": dist,
        "round_dist": dist * 2,
        "path": f"안동→{region}→안동",
    }
    custom = load_custom_destinations()
    custom[name] = entry

    content_str = json.dumps(custom, ensure_ascii=False, indent=2)

    # 로컬 캐시 저장 (같은 세션에서 즉시 반영되도록)
    with open(CUSTOM_DEST_FILE, "w", encoding="utf-8") as f:
        f.write(content_str)

    ok, err = github_commit_file(
        "custom_destinations.json",
        content_str.encode("utf-8"),
        f"도착지 추가: {name}",
    )
    return entry, ok, err


BASE_DESTINATION_DB = {
    "경북소프트웨어마이스터고등학교": {"dist": 47, "round_dist": 94, "path": "안동→의성→안동"},
    "경북에너지기술고등학교": {"dist": 73, "round_dist": 146, "path": "안동→상주→안동"},
    "경북자연과학고등학교": {"dist": 84, "round_dist": 168, "path": "안동→상주→안동"},
    "경북조리과학고등학교": {"dist": 68, "round_dist": 136, "path": "안동→문경→안동"},
    "경북직업훈련교도소": {"dist": 51, "round_dist": 102, "path": "안동→청송→안동"},
    "경북항공고등학교": {"dist": 47, "round_dist": 94, "path": "안동→영주→안동"},
    "경북전문대학교": {"dist": 37, "round_dist": 74, "path": "안동→영주→안동"},
    "문경공업고등학교": {"dist": 48, "round_dist": 96, "path": "안동→문경→안동"},
    "산림조합중앙회임업인종합연수원": {"dist": 65, "round_dist": 130, "path": "안동→청송→안동"},
    "상주공업고등학교": {"dist": 75, "round_dist": 150, "path": "안동→상주→안동"},
    "상주교도소": {"dist": 63, "round_dist": 126, "path": "안동→상주→안동"},
    "상주중장비운전학원": {"dist": 72, "round_dist": 144, "path": "안동→상주→안동"},
    "상지미래경영고등학교": {"dist": 51, "round_dist": 102, "path": "안동→상주→안동"},
    "의성유니텍고등학교": {"dist": 36, "round_dist": 72, "path": "안동→의성→안동"},
    "한국미래농업고등학교": {"dist": 99, "round_dist": 198, "path": "안동→상주→안동"},
    "한국미래산업고등학교": {"dist": 32, "round_dist": 64, "path": "안동→영주→안동"},
    "한국산림과학고등학교": {"dist": 70, "round_dist": 140, "path": "안동→봉화→안동"},
    "한국철도고등학교": {"dist": 33, "round_dist": 66, "path": "안동→영주→안동"},
    "한국철도공사 영주역": {"dist": 31, "round_dist": 62, "path": "안동→영주→안동"},
    "한국펫고등학교": {"dist": 52, "round_dist": 104, "path": "안동→봉화→안동"},
    "한국폴리텍대학 영주캠퍼스": {"dist": 35, "round_dist": 70, "path": "안동→영주→안동"},
    "현대건설중장비직업전문학원": {"dist": 32, "round_dist": 64, "path": "안동→영주→안동"}
}

# 기본 도착지 DB + 사용자가 추가한 도착지 DB 병합
DESTINATION_DB = {**BASE_DESTINATION_DB, **load_custom_destinations()}

# =========================
# 크롤링 및 팝업 제어
# =========================
def close_popups(page):
    try:
        page.keyboard.press("Escape")
    except:
        pass
    try:
        page.evaluate("""
            () => {
                document.querySelectorAll(
                    '[class*="popup"], [id*="popup"], [class*="layer"], .modal, .dim, .overlay'
                ).forEach(el => el.remove());
                document.querySelectorAll('iframe').forEach(el => el.remove());
                document.body.style.overflow = 'auto';
                document.body.style.position = 'static';
                document.querySelectorAll('*').forEach(el => {
                    const style = window.getComputedStyle(el);
                    if (style.position === 'fixed' && parseInt(style.zIndex || 0) > 1000) {
                        el.remove();
                    }
                });
            }
        """)
    except:
        pass

def select_daily_tab(page):
    try:
        page.evaluate("""
            () => {
                const els = Array.from(document.querySelectorAll('a, label, span, button, li'));
                const target = els.find(e => e.textContent.trim() === '일간');
                if (target) { target.click(); return true; }
                return false;
            }
        """)
    except:
        pass

def set_opinet_date(page, date_obj):
    year = str(date_obj.year)
    month = str(date_obj.month)
    day = str(date_obj.day)
    page.evaluate(f"""
        (() => {{
            const target = {{ year: '{year}', month: '{month}', day: '{day}' }};
            function trySet(select, value) {{
                const opt = Array.from(select.options).find(o => o.text.trim() === String(value) || o.value === String(value));
                if (opt) {{
                    select.value = opt.value;
                    select.dispatchEvent(new Event('change', {{ bubbles: true }}));
                    return true;
                }}
                return false;
            }}
            const selects = Array.from(document.querySelectorAll('select')).filter(s => s.offsetParent !== null);
            const yearSelects = selects.filter(s => Array.from(s.options).some(o => o.text.trim() === target.year));
            const monthSelects = selects.filter(s => s.options.length <= 13 && Array.from(s.options).some(o => o.text.trim() === target.month));
            const daySelects = selects.filter(s => s.options.length >= 28 && Array.from(s.options).some(o => o.text.trim() === target.day));
            yearSelects.forEach(s => trySet(s, target.year));
            monthSelects.forEach(s => trySet(s, target.month));
            daySelects.forEach(s => trySet(s, target.day));
        }})()
    """)

def click_query(page):
    for sel in ["a:has-text('조회')", "button:has-text('조회')", "input[value='조회']"]:
        try:
            loc = page.locator(sel)
            if loc.count() > 0:
                loc.first.click(force=True)
                return True
        except:
            pass
    return False

def wait_result_update(page):
    try:
        page.wait_for_load_state("networkidle", timeout=10000)
    except:
        pass
    time.sleep(3)

def navigate_via_menu(page, menu_text, submenu_text):
    close_popups(page)
    try:
        page.get_by_text(menu_text, exact=False).first.hover(timeout=3000)
        page.wait_for_timeout(1000)
    except:
        pass

    close_popups(page)
    try:
        page.get_by_text(submenu_text, exact=False).first.click(force=True, timeout=5000)
        page.wait_for_load_state("domcontentloaded", timeout=15000)
        return True
    except:
        return False


def capture_opinet_print_page(target_date_obj, fuel_type):
    filename = "opinet_capture.png"
    if "LPG" in fuel_type:
        oil_price = 1130
    elif "경유" in fuel_type:
        oil_price = 1510
    else:
        # 휘발유, 하이브리드, 플러그인 하이브리드는 모두 보통휘발유 가격 기준
        oil_price = 1640
    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(
                headless=True,
                executable_path="/usr/bin/chromium",
                args=["--disable-blink-features=AutomationControlled", "--no-sandbox"]
            )
            context = browser.new_context(
                user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                viewport={"width": 1400, "height": 950}
            )
            page = context.new_page()
            if "LPG" in fuel_type:
                # 국내유가통계 > 자동차충전소 메뉴를 거쳐서 진입 (다른 연료와 동일한 방식으로 캡처)
                page.goto("https://www.opinet.co.kr", wait_until="domcontentloaded", timeout=20000)
                page.wait_for_timeout(2000)
                close_popups(page)

                if not navigate_via_menu(page, "국내유가통계", "자동차충전소"):
                    # 메뉴 진입 실패 시 직접 URL로 폴백
                    page.goto(
                        "https://www.opinet.co.kr/user/dopvsavsel/dopVsAvselSelect.do",
                        wait_until="domcontentloaded",
                        timeout=20000
                    )
            else:
                target_url = "https://www.opinet.co.kr/user/dopospdrg/dopOsPdrgSelect.do"
                page.goto(target_url, wait_until="domcontentloaded", timeout=20000)
            page.wait_for_timeout(2000)

            close_popups(page)
            select_daily_tab(page)
            page.wait_for_timeout(800)
            close_popups(page)
            set_opinet_date(page, target_date_obj)
            page.wait_for_timeout(800)

            if not click_query(page):
                raise Exception("조회 버튼 실패")
            wait_result_update(page)

            # 표 구조에 맞춘 유가 파싱 로직
            # 오피넷 표에 날짜가 여러 행(예: 16일, 17일)으로 함께 나오는 경우가 있어,
            # 반드시 조회한 날짜(target_date_obj)와 일치하는 행에서만 값을 가져온다.
            try:
                if "LPG" in fuel_type:
                    target_keyword = "자동차부탄"
                elif "경유" in fuel_type:
                    target_keyword = "자동차용경유"
                else:
                    # 휘발유, 하이브리드, 플러그인 하이브리드는 모두 보통휘발유 가격 기준
                    target_keyword = "보통휘발유"

                y, m, d = target_date_obj.year, target_date_obj.month, target_date_obj.day
                date_candidates = [
                    f"{y}-{m:02d}-{d:02d}",
                    f"{y}.{m:02d}.{d:02d}",
                    f"{y}/{m:02d}/{d:02d}",
                    f"{y}-{m}-{d}",
                    f"{y}.{m}.{d}",
                    f"{m:02d}-{d:02d}",
                    f"{m:02d}.{d:02d}",
                    f"{m}/{d}",
                    f"{m}월{d}일",
                    f"{m:02d}월{d:02d}일",
                    f"{d}일",
                ]
                date_candidates_json = json.dumps(date_candidates, ensure_ascii=False)

                extraction_result = page.evaluate(f"""
                    () => {{
                        const tables = Array.from(document.querySelectorAll('table'));
                        const priceRe = /^[0-9][0-9,]*(\\.[0-9]+)?$/;
                        const dateCandidates = {date_candidates_json};

                        let fallbackPrice = null; // 날짜 매칭 실패 시 사용할 예전 방식(헤더 바로 다음 행) 결과

                        for (const table of tables) {{
                            const rows = Array.from(table.querySelectorAll('tr'));

                            for (let r = 0; r < rows.length; r++) {{
                                const headerCells = Array.from(rows[r].querySelectorAll('th, td'));
                                const colIndex = headerCells.findIndex(c => c.innerText.trim().includes('{target_keyword}'));
                                if (colIndex === -1) continue;

                                for (let dr = r + 1; dr < rows.length; dr++) {{
                                    const dataCells = Array.from(rows[dr].querySelectorAll('th, td'));
                                    if (!dataCells[colIndex]) continue;

                                    const priceText = dataCells[colIndex].innerText.trim();
                                    if (!priceRe.test(priceText)) continue;

                                    // 날짜 매칭 시도: 데이터 행의 맨 앞쪽 셀(들)에서 날짜 문구를 찾는다
                                    const firstCellText = (dataCells[0] ? dataCells[0].innerText.trim() : '');
                                    const rowLabelText = dataCells.slice(0, 2).map(c => c.innerText.trim()).join(' ');
                                    const isDateMatch = dateCandidates.some(dc => firstCellText.includes(dc) || rowLabelText.includes(dc));

                                    if (isDateMatch) {{
                                        return {{ price: priceText, dateMatched: true }};
                                    }}

                                    if (fallbackPrice === null) {{
                                        fallbackPrice = priceText;
                                    }}
                                }}
                            }}
                        }}

                        if (fallbackPrice !== null) {{
                            return {{ price: fallbackPrice, dateMatched: false }};
                        }}
                        return null;
                    }}
                """)

                if extraction_result and extraction_result.get("price"):
                    try:
                        oil_price = float(extraction_result["price"].replace(",", ""))
                    except:
                        pass
                    if not extraction_result.get("dateMatched"):
                        st.warning(
                            f"오피넷 표에서 {target_date_obj.strftime('%m월 %d일')} 날짜에 정확히 일치하는 행을 찾지 못해, "
                            f"표에서 찾은 값({oil_price}원)을 사용했습니다. 캡처본에서 날짜를 다시 확인해 주세요."
                        )
                else:
                    st.warning(f"오피넷에서 실제 유가를 찾지 못해 기본값({oil_price}원)을 사용했습니다. 표 구조가 예상과 다를 수 있습니다.")
            except Exception as ex:
                print(f"세로축 가격 파싱 실패: {ex}")

            print_page = None
            try:
                with context.expect_page(timeout=7000) as pop:
                    try:
                        page.evaluate("chkPrint();")
                    except:
                        page.get_by_text("화면인쇄").first.click(force=True)
                    page.wait_for_timeout(1500)
                print_page = pop.value
            except:
                if len(context.pages) > 1:
                    print_page = context.pages[-1]

            if print_page is None:
                raise Exception("인쇄창 없음")

            print_page.screenshot(path=filename, full_page=True)
            browser.close()
            return filename, oil_price
    except Exception as e:
        st.error(f"오피넷 캡처 중 에러 발생: {e}")
        return filename, oil_price

def find_matched_map_image(dest_name):
    base_dir = os.path.dirname(os.path.abspath(__file__))
    map_dir = os.path.join(base_dir, "map")
    for ext in [".png", ".jpg", ".jpeg", ".PNG", ".JPG", ".JPEG"]:
        target = os.path.join(map_dir, f"{dest_name}{ext}")
        if os.path.exists(target):
            return target

    # 로컬(임시 파일시스템)에 없을 경우, GitHub 저장소에서 다시 받아와 캐시한다.
    # (Streamlit Cloud는 재시작 시 로컬 파일이 초기화되므로, 커밋해둔 원본에서 복구)
    for ext in [".png", ".jpg", ".jpeg"]:
        remote_bytes = github_fetch_file(f"map/{dest_name}{ext}")
        if remote_bytes is not None:
            os.makedirs(map_dir, exist_ok=True)
            target = os.path.join(map_dir, f"{dest_name}{ext}")
            with open(target, "wb") as f_img:
                f_img.write(remote_bytes)
            return target
    return None

# =========================
# DOC 서식 조작용 헬퍼 함수
# =========================
def set_run_font(run, font_name, size_pt, bold=False):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def apply_title_table_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="double" w:sz="12" w:space="0" w:color="000000"/>'
        '<w:left w:val="none"/>'
        '<w:bottom w:val="double" w:sz="12" w:space="0" w:color="000000"/>'
        '<w:right w:val="none"/>'
        '</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def apply_main_table_outer_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="A0A0A0"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="A0A0A0"/>'
        '</w:tblBorders>'
    )
    tblPr.insert_element_before(
        borders, 'w:shd', 'w:tblLayout', 'w:tblCellMar', 'w:tblLook',
        'w:tblCaption', 'w:tblDescription', 'w:tblPrChange',
    )

def remove_cell_margins(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    # [수정 완료] d:type="dxa" 오타를 w:type="dxa"로 수정하여 정상화
    tcMar = parse_xml(
        '<w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:w="0" w:type="dxa"/>'
        '<w:bottom w:w="0" w:type="dxa"/>'
        '<w:left w:w="0" w:type="dxa"/>'
        '<w:right w:w="0" w:type="dxa"/>'
        '</w:tcMar>'
    )
    tcPr.append(tcMar)

def fix_table_width_and_indent(table, width_length):
    """표의 왼쪽 들여쓰기를 0으로 고정하고, 전체 너비를 정확히 고정값으로 지정한다.
    표 스타일(Table Grid 등)마다 기본 들여쓰기/너비 계산 방식이 달라
    제목 표와 본문 표의 좌우 테두리 위치가 어긋나 보이는 문제를 방지하기 위함.
    OOXML의 w:tblPr 자식 요소 순서 규칙(tblW -> jc -> tblCellSpacing -> tblInd
    -> tblBorders -> shd -> tblLayout ...)을 지켜서 삽입해야 워드에서
    "복구" 경고 없이 정상적으로 열린다."""
    tblPr = table._tbl.tblPr
    width_twips = str(int(width_length.twips))

    # python-docx가 표 생성 시 기본으로 넣어두는 <w:tblW w:type="auto"/> 를 제거하고
    # (하나의 tblPr에 tblW가 2개 있으면 스키마 위반으로 워드에서 "복구" 경고가 뜸)
    # 우리가 지정한 고정 너비 값으로 교체한다.
    for old_tblW in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old_tblW)

    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), width_twips)
    tblW.set(qn('w:type'), 'dxa')
    tblPr.insert_element_before(
        tblW, 'w:jc', 'w:tblCellSpacing', 'w:tblInd', 'w:tblBorders',
        'w:shd', 'w:tblLayout', 'w:tblCellMar', 'w:tblLook',
        'w:tblCaption', 'w:tblDescription', 'w:tblPrChange',
    )

    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), '0')
    tblInd.set(qn('w:type'), 'dxa')
    for old_tblInd in tblPr.findall(qn('w:tblInd')):
        tblPr.remove(old_tblInd)
    tblPr.insert_element_before(
        tblInd, 'w:tblBorders', 'w:shd', 'w:tblLayout', 'w:tblCellMar',
        'w:tblLook', 'w:tblCaption', 'w:tblDescription', 'w:tblPrChange',
    )

    tblPr.get_or_add_tblLayout().type = 'fixed'


# =========================
# DOCX 리포트 생성 프로세스
# =========================
def create_docx_report(data_dict, map_image_path, opinet_image_path="opinet_capture.png"):
    doc = Document()

    for section in doc.sections:
        # 용지를 명시적으로 A4로 고정 (기본 템플릿이 Letter/LTR로 되어 있어
        # 인쇄 시 "최적 용지 없음" 오류가 나거나, Letter가 A4보다 짧아
        # 내용이 다음 페이지로 밀려 빈 페이지가 생기는 문제를 방지)
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.orientation = WD_ORIENT.PORTRAIT
        section.top_margin = Mm(10)
        section.bottom_margin = Mm(10)
        section.left_margin = Mm(15)
        section.right_margin = Mm(15)

    total_table_width = Mm(180.0)
    col_widths = [Mm(35.0), Mm(55.0), Mm(35.0), Mm(55.0)]

    # 제목 행 + 데이터 5행 + 이미지 2행 = 총 8행을 "하나의 표"로 만든다.
    # 제목과 본문을 별개의 표 2개로 만들면, 워드 프로그램에 따라 두 표의
    # 실제 렌더링 너비/시작 위치가 미세하게 달라져 좌우 테두리가 어긋나 보일 수 있다.
    # 아예 같은 표로 합쳐버리면 같은 tblGrid(열 구조)를 공유하므로
    # 좌우 경계가 어긋나는 것 자체가 구조적으로 불가능해진다.
    table = doc.add_table(rows=8, cols=4)
    table.style = "Table Grid"
    table.autofit = False
    fix_table_width_and_indent(table, total_table_width)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = col_widths[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 1. 제목 행 (0번 행) - 4개 열을 하나로 병합
    title_cell = table.rows[0].cells[0]
    for c in table.rows[0].cells[1:]:
        title_cell = title_cell.merge(c)
    remove_cell_margins(title_cell)

    set_cell_background(title_cell, "E0E0E0")
    apply_title_table_borders(title_cell)

    title_p = title_cell.paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(0)
    title_p.paragraph_format.line_spacing = 1.0

    title_run = title_p.add_run("시외출장 지출(개인차량) 증빙 내역")
    set_run_font(title_run, "맑은 고딕", 24, bold=True)

    # 2. 데이터 행 (1~5번 행)
    rows_data = [
        ("운행일시", data_dict["date"], "유류비(원)", f"{data_dict['fuel_cost']:,}"),
        ("출장지", data_dict["path"], "통행료", f"{data_dict['toll']:,}"),
        ("거리(km)", f"{data_dict['distance']} km (왕복)", "일 비", f"{data_dict['daily_allowance']:,}"),
        ("연비(km/ℓ)", str(data_dict["efficiency"]), "식 비", f"{data_dict['meal_allowance']:,}"),
        ("유가(원,오피넷기준)", f"{data_dict['oil_price']:,}", "총 계", f"{data_dict['total_cost']:,}"),
    ]

    # 모든 라벨 크기 13pt 및 굵게 적용
    for offset, (l1, v1, l2, v2) in enumerate(rows_data):
        cells = table.rows[offset + 1].cells

        # 1열: 왼쪽 라벨 영역 (13pt, Bold)
        cells[0].text = ""
        p0 = cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_before = Pt(2)
        p0.paragraph_format.space_after = Pt(2)
        set_run_font(p0.add_run(l1), "맑은 고딕", 13, bold=True) 
        set_cell_background(cells[0], "E0E0E0") 

        # 2열: 왼쪽 데이터 영역 (10pt)
        cells[1].text = ""
        p1 = cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after = Pt(2)
        set_run_font(p1.add_run(str(v1)), "맑은 고딕", 10, bold=False)

        # 3열: 오른쪽 라벨 영역 (13pt, Bold)
        cells[2].text = ""
        p2 = cells[2].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after = Pt(2)
        set_run_font(p2.add_run(l2), "맑은 고딕", 13, bold=True) 
        set_cell_background(cells[2], "E0E0E0") 

        # 4열: 오른쪽 데이터 영역 (10pt)
        cells[3].text = ""
        p3 = cells[3].paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_before = Pt(2)
        p3.paragraph_format.space_after = Pt(2)
        set_run_font(p3.add_run(str(v2)), "맑은 고딕", 10, bold=False)

    # 3. 경로 네이버지도 스크린샷 행 처리 (6번 행)
    map_cell = table.rows[6].cells[0]
    for c in table.rows[6].cells[1:]:
        map_cell = map_cell.merge(c)
    remove_cell_margins(map_cell)
    
    map_p = map_cell.paragraphs[0]
    map_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    map_p.paragraph_format.space_before = Pt(2)
    map_p.paragraph_format.space_after = Pt(2)
    
    if map_image_path and os.path.exists(map_image_path):
        map_p.add_run().add_picture(map_image_path, width=Mm(120)) 
    else:
        r = map_p.add_run("경로 네이버지도 스크린샷")
        set_run_font(r, "맑은 고딕", 11, bold=False)

    # 4. 오피넷 스크린샷 행 처리 (7번 행)
    opinet_cell = table.rows[7].cells[0]
    for c in table.rows[7].cells[1:]:
        opinet_cell = opinet_cell.merge(c)
    remove_cell_margins(opinet_cell)
    
    opinet_p = opinet_cell.paragraphs[0]
    opinet_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    opinet_p.paragraph_format.space_before = Pt(2)
    opinet_p.paragraph_format.space_after = Pt(2)
    
    if opinet_image_path and os.path.exists(opinet_image_path):
        opinet_p.add_run().add_picture(opinet_image_path, width=Mm(120))
    else:
        r = opinet_p.add_run("오피넷 스크린샷")
        set_run_font(r, "맑은 고딕", 11, bold=False)

    apply_main_table_outer_borders(table)

    output = "출장지출증빙_보고서.docx"
    doc.save(output)
    return output


# =========================
# Streamlit 웹 UI 실행부
# =========================
st.set_page_config(page_title="출장 지출 증빙", layout="centered")

st.title("시외출장 지출보고서 생성기")

st.markdown("""
<style>
div.stButton > button, 
div.stDownloadButton > button {
    background-color: #191970 !important;
    height: 3.4em !important;
    border-radius: 12px !important;
    border: none !important;
    box-shadow: 0 4px 8px rgba(56, 182, 255, 0.35) !important;
    transition: all 0.2s ease;
}
div.stButton > button p, 
div.stDownloadButton > button p {
    color: white !important;
    font-size: 21px !important;
    font-weight: 600 !important;
}
div.stButton > button:hover, 
div.stDownloadButton > button:hover {
    background-color: #6495ED !important;
    box-shadow: 0 6px 12px rgba(56, 182, 255, 0.45) !important;
}
/* 도착지 드롭다운 목록 맨 아래 '도착지 추가하기' 항목을 다른 색상으로 표시 */
ul[data-testid="stSelectboxVirtualDropdown"] li:last-child,
ul[data-testid="stSelectboxVirtualDropdown"] li:last-child * {
    color: #ff8c00 !important;
    font-weight: 600 !important;
}
</style>
""", unsafe_allow_html=True)


@st.dialog("도착지 추가하기")
def add_destination_dialog():
    st.markdown("네이버 지도를 캡처하고 그 캡처본을 도착지명(ex: 경북소프트웨어마이스터고등학교)으로 저장하여 첨부해 주세요.")

    map_capture = st.file_uploader(
        "네이버 지도 캡처 이미지 첨부",
        type=["png", "jpg", "jpeg"],
        key="new_dest_map_capture",
    )
    new_dest_name = st.text_input(
        "도착지명",
        placeholder="ex: 경북소프트웨어마이스터고등학교",
        key="new_dest_name",
    )
    new_dest_region = st.text_input(
        "도착지 지역명",
        placeholder="ex: 의성, 영주",
        key="new_dest_region",
    )
    new_dest_dist = st.number_input(
        "편도 거리 (km)",
        min_value=0.0,
        step=1.0,
        key="new_dest_dist",
    )

    col_save, col_cancel = st.columns(2)
    with col_save:
        save_clicked = st.button("저장", use_container_width=True)
    with col_cancel:
        cancel_clicked = st.button("취소", use_container_width=True)

    if save_clicked:
        name = new_dest_name.strip()
        region = new_dest_region.strip()

        if not map_capture:
            st.error("네이버 지도 캡처 이미지를 첨부해 주세요.")
        elif not name:
            st.error("도착지명을 입력해 주세요.")
        elif not region:
            st.error("도착지 지역명을 입력해 주세요.")
        elif new_dest_dist <= 0:
            st.error("편도 거리를 입력해 주세요.")
        elif name in BASE_DESTINATION_DB or name in load_custom_destinations():
            st.error("이미 존재하는 도착지명입니다.")
        else:
            ext = (os.path.splitext(map_capture.name)[1] or ".png").lower()
            if ext not in (".png", ".jpg", ".jpeg"):
                ext = ".png"
            image_bytes = map_capture.getbuffer().tobytes()

            # 로컬 캐시 저장 (같은 세션에서 바로 반영되도록)
            os.makedirs(MAP_DIR, exist_ok=True)
            with open(os.path.join(MAP_DIR, f"{name}{ext}"), "wb") as f_img:
                f_img.write(image_bytes)

            with st.spinner("GitHub에 커밋하는 중..."):
                entry, dest_ok, dest_err = save_new_destination(name, region, new_dest_dist)
                img_ok, img_err = github_commit_file(
                    f"map/{name}{ext}",
                    image_bytes,
                    f"도착지 지도 캡처 추가: {name}",
                )

            if dest_ok and img_ok:
                st.session_state.pending_dest_selection = name
                st.success(f"'{name}' 도착지가 추가되어 GitHub에 커밋되었습니다.")
                st.rerun()
            else:
                st.session_state.pending_dest_selection = name
                st.warning(
                    "도착지는 이번 세션에 추가됐지만 GitHub 커밋에 실패했습니다. "
                    "앱이 재시작되면 이 도착지 정보가 사라질 수 있습니다.\n\n"
                    f"{dest_err or ''} {img_err or ''}".strip()
                )
                st.rerun()

    if cancel_clicked:
        fallback = next(iter(DESTINATION_DB), ADD_DEST_OPTION)
        st.session_state.pending_dest_selection = fallback
        st.rerun()


col1, col2 = st.columns(2)
with col1:
    run_date = st.date_input("운행일시", datetime.today(), max_value=datetime.today())
with col2:
    if "pending_dest_selection" in st.session_state:
        st.session_state.dest_selectbox = st.session_state.pop("pending_dest_selection")

    dest_options = list(DESTINATION_DB.keys()) + [ADD_DEST_OPTION]
    dest_selection = st.selectbox("도착지", dest_options, key="dest_selectbox")

if dest_selection == ADD_DEST_OPTION:
    add_destination_dialog()

col3, col4 = st.columns(2)
with col3:
    fuel_selection = st.radio(
        "연료",
        [
            "휘발유 (10.06 km/ℓ)",
            "경유 (10.16 km/ℓ)",
            "LPG (7.87 km/ℓ)",
            "하이브리드 (15.37 km/ℓ)",
            "플러그인 하이브리드 (10.61 km/ℓ)",
        ]
    )
with col4:
    toll_input = st.number_input("통행료", 0, step=100)

col5, col6 = st.columns(2)
with col5:
    daily_fee = st.number_input("일비", 0, step=1000)
with col6:
    meal_fee = st.number_input("식비", 25000, step=1000)

matched_img_file = find_matched_map_image(dest_selection) if dest_selection != ADD_DEST_OPTION else None

st.write("---")

if "report_ready" not in st.session_state:
    st.session_state.report_ready = False

if dest_selection == ADD_DEST_OPTION:
    st.info("도착지를 추가한 뒤, 목록에서 새로 추가된 도착지를 선택해 주세요.")
elif st.button("보고서 생성", use_container_width=True):
    db_info = DESTINATION_DB[dest_selection]
    round_distance = db_info["round_dist"]
    if "휘발유" in fuel_selection and "하이브리드" not in fuel_selection:
        efficiency = 10.06
    elif "경유" in fuel_selection:
        efficiency = 10.16
    elif "LPG" in fuel_selection:
        efficiency = 7.87
    elif "플러그인 하이브리드" in fuel_selection:
        efficiency = 10.61
    elif "하이브리드" in fuel_selection:
        efficiency = 15.37
    else:
        efficiency = 10.06

    with st.spinner("오피넷 조회 중..."):
        opinet_img, oil_price = capture_opinet_print_page(run_date, fuel_selection)

    fuel_cost = int(round_distance / efficiency * oil_price)
    total_cost_raw = fuel_cost + toll_input + daily_fee + meal_fee
    total_cost = int(total_cost_raw // 10) * 10  # 총계만 1원 단위 절사 (예: 1882.72 -> 1880)

    report_data = {
        "date": run_date.strftime("%Y년 %m월 %d일"),
        "path": db_info["path"],
        "distance": round_distance,
        "efficiency": efficiency,
        "oil_price": oil_price,
        "fuel_cost": fuel_cost,
        "toll": toll_input,
        "daily_allowance": daily_fee,
        "meal_allowance": meal_fee,
        "total_cost": total_cost
    }

    file = create_docx_report(report_data, matched_img_file, opinet_img)

    st.session_state.report_ready = True
    st.session_state.report_file = file
    st.session_state.opinet_img = opinet_img
    st.session_state.matched_img_file = matched_img_file

if st.session_state.report_ready:
    st.success("보고서 생성이 완료되었습니다.")
    report_file = st.session_state.report_file
    opinet_img = st.session_state.opinet_img
    matched_img_file = st.session_state.matched_img_file

    with open(report_file, "rb") as f:
        st.download_button(
            "보고서 다운로드",
            f,
            file_name="출장지출증빙_보고서.docx",
            use_container_width=True
        )

    col_a, col_b = st.columns(2)
    with col_a:
        if matched_img_file and os.path.exists(matched_img_file):
            st.image(matched_img_file, caption="네이버지도 경로")
            with open(matched_img_file, "rb") as f_map:
                import base64
                map_bytes = f_map.read()
                b64 = base64.b64encode(map_bytes).decode()
                href = f'data:application/octet-stream;base64,{b64}'
                btn_html = (
                    f'<a href="{href}" download="{os.path.basename(matched_img_file)}" target="_blank" style="text-decoration: none; width: 100%;">'
                    f'  <button style="'
                    f'      width: 100%;'
                    f'      background-color: transparent;'
                    f'      color: rgb(49, 51, 63);'
                    f'      border: 1px solid rgba(49, 51, 63, 0.2);'
                    f'      border-radius: 0.5rem;'
                    f'      padding: 0.4rem 0.75rem;'
                    f'      font-size: 14px;'
                    f'      font-weight: 400;'
                    f'      line-height: 1.6;'
                    f'      cursor: pointer;'
                    f'      text-align: center;'
                    f'      font-family: inherit;'
                    f'  ">경로 캡처 다운받기</button>'
                    f'</a>'
                )
                st.markdown(btn_html, unsafe_allow_html=True)
        else:
            st.warning("map 폴더에서 일치하는 지도 사진을 찾지 못했습니다.")
    with col_b:
        if os.path.exists(opinet_img):
            st.image(opinet_img, caption="오피넷 화면인쇄 증빙")
            with open(opinet_img, "rb") as f_oil:
                import base64
                opinet_bytes = f_oil.read()
                b64 = base64.b64encode(opinet_bytes).decode()
                href = f'data:application/octet-stream;base64,{b64}'
                
                st.markdown(f'''
                    <a href="{href}" download="opinet_capture.png" target="_blank" style="text-decoration: none; width: 100%;">
                        <button style="
                            width: 100%;
                            background-color: transparent;
                            color: rgb(49, 51, 63);
                            border: 1px solid rgba(49, 51, 63, 0.2);
                            border-radius: 0.5rem;
                            padding: 0.4rem 0.75rem;
                            font-size: 14px;
                            font-weight: 400;
                            line-height: 1.6;
                            cursor: pointer;
                            text-align: center;
                            font-family: inherit;
                        ">유가 캡처 다운받기</button>
                    </a>
                ''', unsafe_allow_html=True)
        else:
            st.warning("오피넷 캡처 이미지를 생성하지 못했습니다. 위 에러 메시지를 확인해주세요.")
